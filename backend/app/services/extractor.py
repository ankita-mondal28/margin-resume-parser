"""
Extracts raw text from an uploaded resume file.

Supported inputs:
- .pdf   -> text layer via pdfplumber, falls back to OCR (pdf2image + pytesseract)
            for scanned/image-only PDFs.
- .docx  -> python-docx (paragraphs + tables).
- .png / .jpg / .jpeg -> OCR via pytesseract.
"""

from fastapi import HTTPException

import docx
import pdfplumber
import pytesseract
from pdf2image import convert_from_path
from PIL import Image

MIN_MEANINGFUL_CHARS = 40  # below this, a PDF is treated as "no text layer"


def extract_text(file_path: str, extension: str) -> str:
    if extension == ".pdf":
        return _extract_from_pdf(file_path)
    if extension == ".docx":
        return _extract_from_docx(file_path)
    if extension in (".png", ".jpg", ".jpeg"):
        return _extract_from_image(file_path)
    raise HTTPException(status_code=415, detail=f"Unsupported extension: {extension}")


def _extract_from_pdf(file_path: str) -> str:
    text_chunks = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                text_chunks.append(page_text)
    except Exception as exc:
        raise HTTPException(status_code=422, detail=f"Could not read PDF: {exc}") from exc

    text = "\n".join(text_chunks).strip()

    if len(text) < MIN_MEANINGFUL_CHARS:
        # Likely a scanned/image-only PDF — fall back to OCR.
        text = _ocr_pdf(file_path)

    if len(text.strip()) < MIN_MEANINGFUL_CHARS:
        raise HTTPException(
            status_code=422,
            detail="Couldn't find readable text in this PDF, even with OCR. Try a clearer scan.",
        )
    return text


def _ocr_pdf(file_path: str) -> str:
    try:
        pages = convert_from_path(file_path, dpi=300)
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail=(
                "OCR fallback failed to render the PDF. Make sure 'poppler' is "
                f"installed on the server. ({exc})"
            ),
        ) from exc

    ocr_text = []
    for page_image in pages:
        ocr_text.append(pytesseract.image_to_string(page_image))
    return "\n".join(ocr_text)


def _extract_from_docx(file_path: str) -> str:
    try:
        document = docx.Document(file_path)
    except Exception as exc:
        raise HTTPException(status_code=422, detail=f"Could not read DOCX: {exc}") from exc

    parts = [p.text for p in document.paragraphs if p.text.strip()]

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    text = "\n".join(parts).strip()
    if len(text) < MIN_MEANINGFUL_CHARS:
        raise HTTPException(status_code=422, detail="This DOCX file appears to be empty.")
    return text


def _extract_from_image(file_path: str) -> str:
    try:
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image)
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail=(
                "OCR failed to read this image. Make sure 'tesseract-ocr' is "
                f"installed on the server. ({exc})"
            ),
        ) from exc

    if len(text.strip()) < MIN_MEANINGFUL_CHARS:
        raise HTTPException(
            status_code=422,
            detail="Couldn't find readable text in this image. Try a clearer photo/scan.",
        )
    return text
